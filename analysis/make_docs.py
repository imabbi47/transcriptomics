#!/usr/bin/env python3
"""Render the two-dataset RNA-seq re-analysis report to HTML, .docx and .pdf
from a single content spec."""
from __future__ import annotations

import argparse
import base64
import html as _html

TITLE = ("Re-analysis of two public RNA-seq datasets: genome-wide differential expression, "
         "pathway enrichment, and ciliary gene-set behaviour in a murine tumour model and "
         "human nasal-epithelial cancer")
BYLINE = "Abhishek"
AFFIL = "[Affiliation]  ·  Analysis performed with the open transcriptomics pipeline"

ABSTRACT = (
    "We re-analysed two publicly available RNA-seq datasets end-to-end — genome-wide differential "
    "expression, pathway enrichment, and the behaviour of a curated 683-gene cilia/ciliopathy set. In "
    "syngeneic mouse tumour lines treated with the PDE5 inhibitor sildenafil (GSE334363), 6,553 of 17,426 "
    "tested genes were differentially expressed (FDR<0.05); up-regulated genes were enriched for cholesterol "
    "homeostasis and interferon responses and down-regulated genes for Myc-target and ribosome-biogenesis "
    "programmes, recapitulating the study's cholesterol-trafficking mechanism. In human nasal brushings "
    "comparing cancer to benign tissue (GSE336399), 479 genes were differentially expressed and "
    "predominantly down-regulated (463/479), with strong suppression of inflammatory and interferon "
    "responses. Against these backdrops the cilia set was coordinately down-regulated in human cancer (all "
    "8 significant genes down) but only incidentally and bidirectionally perturbed in the mouse lines (223 "
    "significant). Findings are a descriptive secondary analysis; the two studies differ in species, tissue "
    "and design, and enrichment used over-representation rather than rank-based GSEA."
)

H_PCA = "results/gse336399/pca.png"
H_ENR = "results/gse336399/enrichment_combined.png"
H_VOL = "results/gse336399/volcano_geneset.png"
M_PCA = "results/gse334363/pca.png"
M_ENR = "results/gse334363/enrichment_combined.png"
M_VOL = "results/gse334363/volcano_geneset.png"

BLOCKS = [
    ("h2", "1  Introduction"),
    ("p", "Cilia are microtubule-based organelles: motile cilia line the respiratory epithelium and drive "
          "mucociliary clearance, while primary cilia coordinate signalling central to development and "
          "homeostasis. Their assembly depends on a large, conserved gene programme — intraflagellar "
          "transport, the Bardet-Biedl syndrome (BBS) complex, and small GTPases such as ARL13B — whose "
          "disruption causes the human ciliopathies. Because the nasal and bronchial epithelium is heavily "
          "ciliated, this programme is a candidate readout of epithelial health, and its attenuation has been "
          "linked to loss of differentiated cell states in carcinogenesis."),
    ("p", "We re-analysed two public datasets that were not designed to be compared. For each we first "
          "characterised genome-wide differential expression and pathway enrichment, and then used a defined "
          "683-gene cilia set as a focused lens. Any cross-dataset concordance is treated as "
          "hypothesis-generating rather than confirmatory."),

    ("h2", "2  Methods"),
    ("h3", "2.1  Data sources"),
    ("p", "Two Gene Expression Omnibus (GEO) series were used. GSE336399: 120 human nasal-brushing RNA-seq "
          "samples from a never-smoker lung-cancer field-of-injury study. GSE334363: 18 mouse RNA-seq samples "
          "from three syngeneic tumour cell lines (4T1, MC38, LLC) treated with vehicle or sildenafil (three "
          "replicates each). Raw reads were embargoed at the time of analysis, so the authors' processed "
          "gene-level count matrices were used directly."),
    ("h3", "2.2  Differential expression"),
    ("p", "Counts were analysed with DESeq2 via pyDESeq2. Human samples were restricted to Cancer (n=52) and "
          "Benign (n=40); model ~ sequencing_batch + gender + cancer_status, testing Cancer vs Benign (every "
          "batch contained both groups). Mouse model ~ cell_line + treatment, testing Sildenafil vs Control "
          "with cell line as a blocking factor. Significance: Benjamini-Hochberg adjusted p (FDR) < 0.05."),
    ("h3", "2.3  Pathway enrichment"),
    ("p", "Genes significant at FDR<0.05 were split by direction and submitted to Enrichr (via gseapy) against "
          "MSigDB Hallmark 2020, KEGG 2021 and GO Biological Process 2023; the reported statistic is the "
          "library-adjusted p. This is over-representation analysis on thresholded gene lists, not rank-based "
          "GSEA."),
    ("h3", "2.4  Gene set, matching, and visualisation"),
    ("p", "A user-supplied 683-gene set, dominated by cilia/ciliopathy genes (e.g. BBS1-12, ARL13B, ARL3/6, "
          "B9D1/2, ARMC2/4/9), was intersected with each result table by gene symbol using case-insensitive "
          "matching (human symbols therefore matched identically-spelled mouse orthologs; divergently named "
          "orthologs were not mapped). Sample structure was assessed by principal-component analysis (PCA) of "
          "log-normalised counts, and results were visualised with volcano plots. Software: Python 3.12, "
          "pyDESeq2, gseapy, pandas, scikit-learn, matplotlib."),

    ("h2", "3  Results"),

    ("h3", "3.1  Human nasal epithelium (GSE336399)"),
    ("p", "Principal-component analysis (Figure 1) summarises sample structure; the cancer-versus-benign "
          "contrast is a subtle field-of-injury effect rather than the dominant axis of variance. Genome-wide, "
          "479 genes were significant and overwhelmingly down-regulated (463 down vs 16 up). The strongest "
          "up-regulated gene was the desmosomal DSG4 (log2FC +14.6); down-regulated genes included "
          "immunoglobulin loci (IGHV3-23, IGHD) and P2RY14 (Table 1)."),
    ("fig", H_PCA, "Figure 1. PCA of GSE336399 samples (log-normalised counts). Points are coloured by "
                   "cancer status; the field-of-injury contrast is subtle relative to overall variance."),
    ("table", "Table 1. Representative genome-wide differentially expressed genes, human (by adjusted p).",
     ["Gene", "log2FC", "Adj. p"],
     [["DSG4", "+14.61", "2.7e-60"], ["KRT32", "+8.13", "3.3e-5"], ["KRTAP4-1", "+8.22", "1.3e-4"],
      ["IGHV3-23", "-7.52", "3.5e-4"], ["P2RY14", "-2.80", "6.5e-4"], ["IGHD", "-6.05", "1.4e-3"]]),
    ("p", "Enrichment of the down-regulated genes was dominated by immune signalling — Inflammatory Response "
          "(Hallmark, FDR 1.2e-20) and Interferon Gamma Response (FDR 2.6e-15) — indicating suppression of "
          "inflammatory/interferon programmes in the cancerous field, consistent with the source study "
          "(Figure 2). Up-regulated enrichment was weak and borderline (Apoptosis, oxidative phosphorylation)."),
    ("fig", H_ENR, "Figure 2. Pathway enrichment for GSE336399 (A, up-regulated; B, down-regulated genes). "
                   "Down-regulated genes are strongly enriched for inflammatory and interferon responses."),
    ("p", "Within this contraction, the cilia set behaved coherently: of 638 present (589 testable) cilia "
          "genes, 8 were significant and all 8 were down-regulated in cancer (Table 2, Figure 3), with no "
          "significant up-regulated cilia gene."),
    ("fig", H_VOL, "Figure 3. Volcano plot for GSE336399 with the cilia set overlaid (grey, all genes; orange, "
                   "cilia-set genes not significant; blue, significantly down-regulated). y-axis scaled to the set."),
    ("table", "Table 2. All significant cilia-set genes, human (FDR<0.05) — all down-regulated.",
     ["Gene", "log2FC", "Adj. p"],
     [["P2RY14", "-2.80", "6.5e-4"], ["LRRK2", "-1.84", "0.013"], ["RIPOR2", "-2.13", "0.014"],
      ["SNX10", "-1.96", "0.018"], ["TGFB1", "-1.29", "0.025"], ["RCSD1", "-1.51", "0.032"],
      ["TOPORS", "-0.57", "0.039"], ["OPN1MW2", "-2.80", "0.047"]]),

    ("h3", "3.2  Mouse tumour model (GSE334363)"),
    ("p", "PCA (Figure 4) is dominated by cell-line identity — 4T1, MC38 and LLC form three clusters — which "
          "is why treatment was modelled within cell line. Sildenafil produced a broad response: 6,553 of "
          "17,426 tested genes were significant (3,397 up, 3,156 down). Leading genes were "
          "lipid/cholesterol-metabolism related (Rnf145, Sc5d, Fads2) alongside Tuba1a and Nup210 (Table 3)."),
    ("fig", M_PCA, "Figure 4. PCA of GSE334363 samples; the three tumour cell lines dominate sample-to-sample "
                   "variance, with treatment a within-line effect."),
    ("table", "Table 3. Representative genome-wide differentially expressed genes, mouse (by adjusted p).",
     ["Gene", "log2FC", "Adj. p"],
     [["Rnf145", "+1.22", "1.0e-144"], ["Sc5d", "+1.06", "1.2e-95"], ["Itgb5", "+0.51", "7.7e-93"],
      ["Nup210", "-0.72", "4.0e-89"], ["Fads2", "+1.03", "1.1e-88"], ["Tuba1a", "+1.02", "1.2e-87"]]),
    ("p", "Enrichment recapitulated the study's mechanism: up-regulated genes were enriched for Cholesterol "
          "Homeostasis and interferon (alpha/gamma) responses, while down-regulated genes were enriched for "
          "Myc targets and ribosome biogenesis (Figure 5) — together consistent with sildenafil's reported "
          "disruption of NPC1-mediated cholesterol trafficking and an anti-proliferative shift."),
    ("fig", M_ENR, "Figure 5. Pathway enrichment for GSE334363 (A, up-regulated — cholesterol/interferon; "
                   "B, down-regulated — Myc targets, ribosome biogenesis)."),
    ("p", "The cilia set here comprised 560 present (513 testable) genes, of which 223 were significant, split "
          "128 up- and 95 down-regulated (Table 4, Figure 6). Effects were modest but highly significant, led "
          "by Tuba1a; most leading hits are general cytoskeletal/trafficking genes."),
    ("fig", M_VOL, "Figure 6. Volcano plot for GSE334363 with the cilia set overlaid (red/blue, significantly "
                   "up/down-regulated cilia-set genes). The response is bidirectional."),
    ("table", "Table 4. Representative top cilia-set genes, mouse (by adjusted p).",
     ["Gene", "log2FC", "Adj. p"],
     [["Tuba1a", "+1.02", "1.2e-87"], ["Snap29", "+0.51", "6.3e-41"], ["Flna", "-0.21", "2.7e-19"],
      ["Rab11a", "+0.28", "9.1e-19"], ["Cfap298", "-0.36", "8.8e-16"], ["Cdh23", "-0.89", "6.6e-15"]]),

    ("h3", "3.3  Cross-dataset comparison"),
    ("p", "Only one cilia-set gene, TGFB1, was significant in both datasets, and in opposite directions "
          "(mouse log2FC +0.19, FDR 9e-4; human log2FC -1.29, FDR 0.025). As a pleiotropic signalling cytokine "
          "rather than a core structural cilia gene, this overlap is unlikely to reflect shared ciliary biology."),

    ("h2", "4  Discussion"),
    ("p", "Encouragingly, the genome-wide re-analysis recovers each study's headline biology. In the mouse "
          "model, up-regulated cholesterol-homeostasis genes and top hits (Rnf145, Sc5d, Fads2) align with "
          "sildenafil's reported disruption of NPC1-mediated cholesterol trafficking, while down-regulated "
          "Myc-target and ribosome-biogenesis programmes indicate an anti-proliferative shift. In the human "
          "cohort, the predominantly down-regulated signature and suppressed inflammatory/interferon responses "
          "match the reported immune suppression in the malignant field. This concordance supports the "
          "validity of a processed-count re-analysis."),
    ("p", "Against the human background, the ciliary programme is coordinately down-regulated (all 8 "
          "significant cilia genes down), consistent with loss of ciliated, terminally differentiated "
          "epithelial states during carcinogenesis, and nominates specific genes (P2RY14, RIPOR2, SNX10) for "
          "follow-up. In the mouse lines the 223 significant cilia genes are best read as incidental: they sit "
          "within a 6,553-gene response, are bidirectional, and are led by general cytoskeletal/trafficking "
          "genes — and tumour cell lines largely lack motile cilia. The datasets are therefore not directly "
          "comparable, and the single shared hit (TGFB1, opposite directions) reinforces that caution."),
    ("note", "Limitations. (1) Secondary analysis of processed count matrices; embargoed raw reads precluded "
             "read-level QC and independent quantification. (2) Human PCA is dominated by non-contrast variance, "
             "so the field-of-injury effect is subtle. (3) Enrichment is over-representation on thresholded "
             "gene lists, not rank-based GSEA, and no competitive set-level test was applied to the cilia set. "
             "(4) Cross-species matching was symbol-based; divergently named orthologs are missed. (5) The cilia "
             "set is user-defined and mixes motile- and primary-cilia genes. (6) The two studies differ in "
             "species, tissue, platform and design."),
    ("h3", "Future directions"),
    ("p", "A rank-based GSEA (or a Wilcoxon test of cilia-set versus background log-fold changes) would put a "
          "p-value on the human ciliary down-shift; proper mouse-human ortholog mapping would complete the "
          "matching; and ciliated-cell-resolved single-cell data would sharpen the biological interpretation."),

    ("h2", "Data and code availability"),
    ("p", "Datasets: GEO GSE336399 (human nasal epithelium) and GSE334363 (mouse tumour lines). Analysis code: "
          "the open-source transcriptomics pipeline (github.com/imabbi47/transcriptomics)."),

    ("h2", "References"),
    ("refs", [
        "Love MI, Huber W, Anders S. Moderated estimation of fold change and dispersion for RNA-seq data with DESeq2. Genome Biology 2014;15:550.",
        "Muzellec B, Telenczuk M, Cabeli V, Andreux M. PyDESeq2: a Python package for bulk RNA-seq differential expression analysis. Bioinformatics 2023.",
        "Kuleshov MV, et al. Enrichr: a comprehensive gene set enrichment analysis web server 2016 update. Nucleic Acids Research 2016;44:W90-W97.",
        "Reiter JF, Leroux MR. Genes and molecular pathways underpinning ciliopathies. Nature Reviews Molecular Cell Biology 2017;18:533-547.",
        "GSE336399, Gene Expression Omnibus — human nasal-epithelial field-of-injury RNA-seq.",
        "GSE334363, Gene Expression Omnibus — sildenafil-treated murine tumour RNA-seq.",
    ]),
]

CSS = """
*{box-sizing:border-box} body{margin:0}
.paper{background:#fff;color:#1a1d21;max-width:800px;margin:0 auto;padding:34px 26px 70px;
  font-family:Georgia,"Times New Roman",serif;font-size:16.5px;line-height:1.62}
.paper h1{font-family:-apple-system,"Segoe UI",Roboto,sans-serif;font-size:1.6rem;line-height:1.25;
  text-align:center;margin:.2em 0 .4em;text-wrap:balance}
.byline{text-align:center;color:#5b6570;font-family:-apple-system,"Segoe UI",sans-serif;font-size:.95rem}
.affil{text-align:center;color:#5b6570;font-style:italic;font-size:.85rem;margin-bottom:20px}
.tag{text-align:center;font-family:-apple-system,"Segoe UI",sans-serif;font-size:.72rem;letter-spacing:.08em;
  text-transform:uppercase;color:#8a939d;margin-bottom:18px}
.abstract{background:#f5f7f9;border:1px solid #d7dce1;border-radius:8px;padding:16px 20px;font-size:.95rem;margin:0 0 26px}
.paper h2{font-family:-apple-system,"Segoe UI",Roboto,sans-serif;font-size:1.12rem;margin:30px 0 8px;
  padding-bottom:5px;border-bottom:2px solid #1f3554;color:#1f3554}
.paper h3{font-family:-apple-system,"Segoe UI",sans-serif;font-size:.98rem;margin:18px 0 4px;color:#33424f}
.paper p{margin:.5em 0;text-align:justify}
figure{margin:20px 0;text-align:center} figure img{max-width:100%;border:1px solid #d7dce1;border-radius:6px}
figcaption{font-family:-apple-system,"Segoe UI",sans-serif;font-size:.8rem;color:#5b6570;margin-top:7px;text-align:left}
.wrap-x{overflow-x:auto;margin:16px 0}
table{border-collapse:collapse;width:100%;font-family:-apple-system,"Segoe UI",sans-serif;font-size:.82rem}
caption{font-family:-apple-system,"Segoe UI",sans-serif;font-size:.8rem;color:#5b6570;text-align:left;margin-bottom:6px}
th,td{border-bottom:1px solid #d7dce1;padding:6px 10px;text-align:left}
th{color:#1f3554;border-bottom:2px solid #1f3554} td:nth-child(n+2){text-align:right;font-variant-numeric:tabular-nums}
.refs{font-size:.82rem;color:#33424f} .refs ol{padding-left:20px} .refs li{margin:5px 0}
.note{background:#fff7ed;border-left:3px solid #e0872b;padding:10px 14px;font-size:.86rem;margin:16px 0;border-radius:0 6px 6px 0}
"""


def render_html(path):
    def uri(p):
        return "data:image/png;base64," + base64.b64encode(open(p, "rb").read()).decode()

    e = _html.escape
    out = [f"<title>RNA-seq re-analysis report</title>", f"<style>{CSS}</style>",
           '<article class="paper">', '<div class="tag">Secondary re-analysis · Transcriptomics</div>',
           f"<h1>{e(TITLE)}</h1>", f'<div class="byline">{e(BYLINE)}</div>',
           f'<div class="affil">{e(AFFIL)}</div>', f'<div class="abstract"><b>Abstract.</b> {e(ABSTRACT)}</div>']
    for b in BLOCKS:
        k = b[0]
        if k == "h2":
            out.append(f"<h2>{e(b[1])}</h2>")
        elif k == "h3":
            out.append(f"<h3>{e(b[1])}</h3>")
        elif k == "p":
            out.append(f"<p>{e(b[1])}</p>")
        elif k == "note":
            out.append(f'<div class="note">{e(b[1])}</div>')
        elif k == "fig":
            out.append(f'<figure><img src="{uri(b[1])}" alt=""><figcaption>{e(b[2])}</figcaption></figure>')
        elif k == "table":
            _, capt, headers, rows = b
            th = "".join(f"<th>{e(h)}</th>" for h in headers)
            body = "".join("<tr>" + "".join(f"<td>{e(str(c))}</td>" for c in r) + "</tr>" for r in rows)
            out.append(f'<div class="wrap-x"><table><caption>{e(capt)}</caption>'
                       f"<thead><tr>{th}</tr></thead><tbody>{body}</tbody></table></div>")
        elif k == "refs":
            out.append('<div class="refs"><ol>' + "".join(f"<li>{e(r)}</li>" for r in b[1]) + "</ol></div>")
    out.append("</article>")
    open(path, "w", encoding="utf-8").write("\n".join(out))


def render_docx(path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    doc.add_heading(TITLE, level=0)
    for txt, sz, ital in [(BYLINE, 11, False), (AFFIL, 9, True)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.font.size = Pt(sz)
        r.italic = ital
    ab = doc.add_paragraph()
    ab.add_run("Abstract. ").bold = True
    ab.add_run(ABSTRACT)
    for b in BLOCKS:
        k = b[0]
        if k == "h2":
            doc.add_heading(b[1], level=1)
        elif k == "h3":
            doc.add_heading(b[1], level=2)
        elif k in ("p", "note"):
            p = doc.add_paragraph()
            run = p.add_run(b[1])
            run.italic = (k == "note")
        elif k == "fig":
            doc.add_picture(b[1], width=Inches(5.7))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c = doc.add_paragraph().add_run(b[2])
            c.italic = True
            c.font.size = Pt(8.5)
        elif k == "table":
            _, capt, headers, rows = b
            c = doc.add_paragraph().add_run(capt)
            c.italic = True
            c.font.size = Pt(9)
            tbl = doc.add_table(rows=1, cols=len(headers))
            tbl.style = "Light Grid Accent 1"
            for i, h in enumerate(headers):
                tbl.rows[0].cells[i].text = h
            for row in rows:
                cells = tbl.add_row().cells
                for i, v in enumerate(row):
                    cells[i].text = str(v)
        elif k == "refs":
            for i, ref in enumerate(b[1], 1):
                p = doc.add_paragraph(f"{i}. {ref}")
                p.runs[0].font.size = Pt(9)
    doc.save(path)


def render_pdf(path):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle)

    def esc(s):
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    ss = getSampleStyleSheet()
    navy = colors.HexColor("#1f3554")
    muted = colors.HexColor("#5b6570")
    h1 = ParagraphStyle("t", parent=ss["Title"], fontSize=14, leading=18, alignment=TA_CENTER)
    center = ParagraphStyle("c", parent=ss["Normal"], alignment=TA_CENTER, fontSize=9, textColor=muted)
    body = ParagraphStyle("b", parent=ss["Normal"], alignment=TA_JUSTIFY, fontSize=10, leading=14, spaceAfter=5)
    h2 = ParagraphStyle("h2", parent=ss["Heading2"], textColor=navy, fontSize=12, spaceBefore=10)
    h3 = ParagraphStyle("h3", parent=ss["Heading3"], fontSize=10.5, spaceBefore=6)
    cap = ParagraphStyle("cap", parent=ss["Normal"], fontSize=8, textColor=muted, spaceAfter=6)
    abst = ParagraphStyle("ab", parent=body, fontSize=9, leading=12, backColor=colors.HexColor("#f5f7f9"), borderPadding=8)
    note = ParagraphStyle("nt", parent=body, fontSize=9, backColor=colors.HexColor("#fff7ed"), borderPadding=6)
    refst = ParagraphStyle("rf", parent=body, fontSize=8.5, leading=11, spaceAfter=2)

    usable = A4[0] - 2 * inch
    story = [Paragraph(esc(TITLE), h1), Paragraph(esc(BYLINE), center), Paragraph(esc(AFFIL), center),
             Spacer(1, 10), Paragraph("<b>Abstract.</b> " + esc(ABSTRACT), abst), Spacer(1, 8)]
    for b in BLOCKS:
        k = b[0]
        if k == "h2":
            story.append(Paragraph(esc(b[1]), h2))
        elif k == "h3":
            story.append(Paragraph(esc(b[1]), h3))
        elif k == "p":
            story.append(Paragraph(esc(b[1]), body))
        elif k == "note":
            story.append(Paragraph("<b>" + esc(b[1]) + "</b>", note))
        elif k == "fig":
            img = Image(b[1])
            img.drawWidth = usable
            img.drawHeight = usable * img.imageHeight / img.imageWidth
            story += [Spacer(1, 4), img, Paragraph(esc(b[2]), cap)]
        elif k == "table":
            _, capt, headers, rows = b
            story += [Spacer(1, 4), Paragraph(esc(capt), cap)]
            tbl = Table([headers] + rows, hAlign="LEFT",
                        colWidths=[usable * 0.5, usable * 0.25, usable * 0.25])
            tbl.setStyle(TableStyle([
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("TEXTCOLOR", (0, 0), (-1, 0), navy),
                ("LINEBELOW", (0, 0), (-1, 0), 1, navy),
                ("LINEBELOW", (0, 1), (-1, -1), 0.3, colors.HexColor("#d7dce1")),
                ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
                ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(tbl)
        elif k == "refs":
            for i, ref in enumerate(b[1], 1):
                story.append(Paragraph(f"{i}. {esc(ref)}", refst))
    SimpleDocTemplate(path, pagesize=A4, topMargin=0.9 * inch, bottomMargin=0.9 * inch,
                      leftMargin=inch, rightMargin=inch, title="RNA-seq re-analysis report").build(story)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--html")
    ap.add_argument("--docx")
    ap.add_argument("--pdf")
    args = ap.parse_args()
    if args.html:
        render_html(args.html)
        print(f"[html] wrote {args.html}")
    if args.docx:
        render_docx(args.docx)
        print(f"[docx] wrote {args.docx}")
    if args.pdf:
        render_pdf(args.pdf)
        print(f"[pdf]  wrote {args.pdf}")


if __name__ == "__main__":
    main()
