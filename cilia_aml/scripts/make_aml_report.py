#!/usr/bin/env python3
"""Render the Cilia+AML (miR-551b) re-analysis report to HTML, .docx and .pdf."""
from __future__ import annotations
import argparse, base64, html as _html

TITLE = ("Transcriptomic effect of miR-551b inhibition in acute myeloid leukaemia cells, and the "
         "behaviour of a ciliary/centrosomal gene set and a curated interaction network")
BYLINE = "Abhishek"
AFFIL = "[Affiliation]  ·  Analysis performed end-to-end with the open transcriptomics pipeline"

ABSTRACT = (
    "Cilia and centrosomal genes have been proposed as candidate markers in cancer, and miR-551b is an "
    "oncogenic microRNA in several solid tumours. We tested whether inhibiting miR-551b reshapes the "
    "transcriptome of acute myeloid leukaemia (AML) cells and, in particular, whether it perturbs a curated "
    "683-gene cilia/ciliopathy set or a 25-gene STRING-style interaction network (dopaminergic, innate-immune, "
    "autophagy, iron and ciliary genes). Eight single-end RNA-seq runs (BioProject PRJNA1513000; primary AML "
    "blasts and the KG1a cell line, each treated with LNA or ZIP anti-miR against miR-551b versus a control "
    "oligonucleotide) were processed from raw reads: multi-connection download from the SRA AWS Open-Data "
    "mirror, Salmon selective-alignment quantification against GENCODE v44 (72-83% mapping), and DESeq2 "
    "differential expression. Using a paired, block design (adjusting for cell type and antisense chemistry), "
    "only three genes were significant (FDR<0.05; all down-regulated: TCF7, IGKC and one unannotated locus). "
    "The cilia set showed no significant gene (0/563) and was in fact less variable than the transcriptome "
    "background (mean |log2FC| 0.185 vs 0.321); the interaction network showed no significant gene (0/17 "
    "expressed; the eight dopaminergic members were not detectably expressed in AML). Analysed separately, "
    "primary AML samples yielded 33 significant genes dominated by T- and B-lymphocyte markers (TRBC1, ZAP70, "
    "LCK, IGKC, IGHM) - a signature attributable to differing immune-cell content between samples rather than "
    "to miR-551b - while the clean KG1a line yielded essentially none. We conclude that, in this dataset, "
    "miR-551b inhibition produces a minimal and non-reproducible transcriptional footprint in AML cells and "
    "does not engage ciliary/centrosomal genes. This is a negative result and should be read against a clear "
    "power limitation: the design provides only one biological unit per unique condition, with the two "
    "chemistries used as pseudo-replicates.")

F_PCA = "cilia_aml/figures/pca.png"
F_VALL = "cilia_aml/figures/volcano_all.png"
F_VCIL = "cilia_aml/figures/volcano_cilia.png"
F_VNET = "cilia_aml/figures/volcano_network.png"

BLOCKS = [
    ("h2", "1  Introduction"),
    ("p", "Primary cilia and the centrosomal/basal-body apparatus coordinate signalling (Hedgehog, Wnt, "
          "PDGFR) and cell-cycle progression, and their gene programme has been proposed as a candidate axis "
          "in tumour biology. Whether this programme is relevant to acute myeloid leukaemia (AML) is far from "
          "obvious: haematopoietic cells are largely non-ciliated, so any signal is more plausibly centrosomal "
          "than motile-ciliary. miR-551b, in turn, is an oncogenic microRNA amplified in ovarian and gastric "
          "cancers; antisense inhibition of miR-551b has been used to probe its downstream programme."),
    ("p", "Here we ask a focused question: does inhibiting miR-551b in AML cells change the expression of "
          "(i) a curated 683-gene cilia/ciliopathy set and (ii) a 25-gene interaction network (assembled from a "
          "STRING-style neighbourhood spanning dopaminergic signalling, innate immunity, autophagy, iron "
          "handling and two core ciliary genes, IFT88 and ARL13B)? We processed the data end-to-end from raw "
          "reads and treat any result as hypothesis-level, given the study's limited replication."),

    ("h2", "2  Methods"),
    ("h3", "2.1  Data"),
    ("p", "Eight single-end Illumina HiSeq 2000 RNA-seq runs were obtained from NCBI BioProject PRJNA1513000 "
          "(runs SRR40166669-SRR40166676). The design is a balanced 2x2x2: cell context (primary AML blasts "
          "vs the KG1a AML cell line) x antisense chemistry (LNA vs ZIP) x treatment (anti-miR-551b vs control "
          "oligonucleotide). No processed count matrix was deposited (no linked GEO series), so raw reads were "
          "used. Because the archive throttles single-stream transfers, reads were fetched at full depth "
          "(13-25 million reads/run) from the SRA AWS Open-Data mirror with a 16-connection downloader."),
    ("h3", "2.2  Quantification"),
    ("p", "Reads were quantified with Salmon 2.5.1 in selective-alignment mode against a GENCODE v44 "
          "transcriptome index (251,955 transcripts); transcript estimates were summed to gene level "
          "(tximport-style). Mapping rates were 71.8-72.7% for primary-AML samples and 80.7-82.9% for KG1a, "
          "consistent with cleaner cell-line input. The gene x sample matrix contained 60,883 genes; 25,994 "
          "passed a total-count >= 10 filter and were tested."),
    ("h3", "2.3  Differential expression and gene sets"),
    ("p", "Differential expression used DESeq2 (pyDESeq2). The primary model was a paired/blocked design, "
          "~ cell_type + chemistry + treatment, testing anti-miR-551b vs control and thereby removing both the "
          "cell-context and chemistry axes; secondary per-cell-type models (~ chemistry + treatment; n=2 vs 2) "
          "were also fit. Significance was Benjamini-Hochberg FDR < 0.05. Two gene sets were intersected with "
          "the results by gene symbol: a user-supplied 683-gene cilia/ciliopathy set and the 25-gene "
          "interaction network. Software: Python 3.12, Salmon, pyDESeq2, gseapy, pandas, scikit-learn, "
          "matplotlib."),

    ("h2", "3  Results"),
    ("h3", "3.1  Sample structure"),
    ("p", "Principal-component analysis (Figure 1) is dominated by cell context: primary-AML and KG1a samples "
          "separate completely, which is why cell type is a blocking term in the model. Within each block the "
          "anti-miR-551b and control samples do not form a clear separate axis."),
    ("fig", F_PCA, "Figure 1. PCA of the eight samples (DESeq2-normalised, log-scaled counts). Colour = "
                   "treatment (anti-miR-551b vs control); the dominant split is cell context (primary AML vs KG1a)."),
    ("h3", "3.2  Genome-wide differential expression (paired model)"),
    ("p", "Adjusting for cell type and chemistry, only three genes were significant at FDR<0.05, all "
          "down-regulated upon miR-551b inhibition (Table 1, Figure 2). The strongest annotated hits were TCF7 "
          "(a Wnt/T-cell transcription factor) and IGKC (an immunoglobulin constant region); the third is an "
          "unannotated locus. This is a near-null genome-wide response."),
    ("fig", F_VALL, "Figure 2. Volcano plot, anti-miR-551b vs control (paired model). Red points are the three "
                    "FDR-significant genes; the bulk of the transcriptome is unchanged."),
    ("table", "Table 1. All genome-wide significant genes (paired model, FDR<0.05).",
     ["Gene", "log2FC", "Adj. p"],
     [["ENSG00000257379", "-8.46", "1.6e-3"], ["IGKC", "-1.71", "1.5e-2"], ["TCF7", "-2.16", "1.7e-2"]]),
    ("h3", "3.3  The cilia/centrosomal set does not respond"),
    ("p", "Of the 683-gene cilia set, 563 genes were expressed and testable; none reached FDR<0.05 (Figure 3). "
          "Beyond simple non-significance, cilia-set genes were less variable than the transcriptome as a whole "
          "(mean absolute log2 fold-change 0.185 vs 0.321 background) - i.e. they are among the more stable "
          "genes under miR-551b inhibition. The most-shifted cilia genes by nominal p (Table 2) are unremarkable "
          "and do not survive multiple-testing correction."),
    ("fig", F_VCIL, "Figure 3. Volcano plot with the 683-gene cilia set overlaid (orange). No cilia gene is "
                    "significant; the set sits within the unperturbed central cloud."),
    ("table", "Table 2. Most-shifted cilia-set genes by nominal p-value (none FDR-significant).",
     ["Gene", "log2FC", "Nominal p", "Adj. p"],
     [["SYNE1", "-0.52", "0.012", "1.00"], ["CENPF", "+0.59", "0.035", "1.00"],
      ["HAVCR1", "-1.11", "0.050", "1.00"], ["RP1", "+2.01", "0.062", "1.00"],
      ["KIF14", "+0.58", "0.063", "1.00"], ["GSK3B", "-0.28", "0.082", "1.00"]]),
    ("h3", "3.4  The interaction network does not respond"),
    ("p", "Seventeen of the 25 network genes were expressed in AML; none was significant (Table 3, Figure 4). "
          "The eight non-detected members were exactly the dopaminergic module (SLC6A3, DBH, ANKK1, DRD5, "
          "PPP1R1B, DRD1, DRD2) plus IL13 - neurotransmitter genes not expressed in myeloid cells, as expected. "
          "The two ciliary members that anchored the network, IFT88 (-0.04) and ARL13B (+0.21), were flat, as "
          "were the innate-immune (TLR2/TLR4, NFKB1, JAK2), autophagy (ATG7, LAMP2, HSPA8) and iron (FTL, FTH1) "
          "members."),
    ("fig", F_VNET, "Figure 4. Volcano plot with the 25-gene interaction network overlaid (orange). All "
                    "expressed members are non-significant."),
    ("table", "Table 3. Interaction-network genes, ranked by |log2FC| (paired model; all FDR non-significant).",
     ["Gene", "log2FC", "Adj. p"],
     [["TLR4", "+0.47", "~1"], ["GYPE", "+0.23", "~1"], ["ARL13B", "+0.21", "~1"],
      ["TRAF6", "-0.15", "~1"], ["HSPA8", "+0.12", "~1"], ["IFT88", "-0.04", "~1"]]),
    ("h3", "3.5  Per-cell-type analysis reveals a composition confound"),
    ("p", "Analysed separately, the two contexts diverge sharply. Primary AML gave 33 significant genes (8 up, "
          "25 down), but the leading hits are overwhelmingly lymphocyte markers - TRBC1, ZAP70, LCK (T-cell), "
          "IGKC, IGHM (B-cell), plus RORA and TCF7. Because these are patient samples with variable normal "
          "immune-cell content, and each unique condition is n=1, this signature most plausibly reflects "
          "differing T-/B-cell contamination between the anti-miR and control samples rather than a miR-551b "
          "programme. The clean KG1a line gave only three hits, dominated by read-through/unannotated loci - "
          "i.e. essentially no reproducible effect. The pooled paired model (Section 3.2) is therefore the "
          "trustworthy read-out, and it is near-null."),

    ("h2", "4  Discussion"),
    ("p", "Across a genome-wide paired analysis, a clean-cell-line analysis, and two curated gene sets, "
          "inhibition of miR-551b left the AML transcriptome largely unchanged, and specifically did not "
          "engage ciliary/centrosomal genes - which were, if anything, more stable than average. For the "
          "cilia-in-AML hypothesis this is a clear negative: miR-551b is not an upstream regulator of the "
          "ciliary programme in these cells, and the curated interaction network is likewise unresponsive."),
    ("p", "The one apparently strong signal - 33 genes in primary AML - is best read as a cautionary tale about "
          "cell composition in bulk RNA-seq of patient material: an immune-cell-content difference between "
          "single samples masquerades as a treatment effect. It is exactly the kind of result that replication "
          "and deconvolution, not deeper sequencing, would resolve."),
    ("note", "Limitations. (1) The dominant limitation is replication: the 2x2x2 design provides one biological "
             "unit per unique condition, with the LNA and ZIP chemistries used as pseudo-replicates; power to "
             "detect modest effects is low. (2) miRNA inhibition often acts translationally, so a small mRNA "
             "footprint does not exclude a proteomic effect. (3) Primary-AML samples carry variable normal-cell "
             "content, confounding the per-sample contrast. (4) Quantification was alignment-free "
             "(selective alignment) rather than genome alignment. (5) Gene-set matching was by symbol; the cilia "
             "set mixes motile- and primary-cilia genes and is not myeloid-tailored. (6) No competitive set-level "
             "enrichment test was applied (with 0 hits, over-representation is uninformative)."),
    ("h3", "Future directions"),
    ("p", "A replicated design (biological triplicates per condition), computational immune-cell deconvolution "
          "of the primary-AML samples, and a miR-551b target-site-aware analysis (e.g. testing predicted "
          "targets as a set) would convert this negative screen into a definitive test. If the centrosomal "
          "angle is of interest, a ciliated/centrosome-resolved model system would be more informative than "
          "bulk myeloid RNA-seq."),

    ("h2", "Data and code availability"),
    ("p", "Data: NCBI BioProject PRJNA1513000 (runs SRR40166669-SRR40166676). Reference: GENCODE v44. "
          "Analysis code: the open-source transcriptomics pipeline (github.com/imabbi47/transcriptomics); "
          "all scripts, designs and result tables for this analysis are under cilia_aml/."),
    ("h2", "References"),
    ("refs", [
        "Love MI, Huber W, Anders S. Moderated estimation of fold change and dispersion for RNA-seq data with DESeq2. Genome Biology 2014;15:550.",
        "Patro R, Duggal G, Love MI, Irizarry RA, Kingsford C. Salmon provides fast and bias-aware quantification of transcript expression. Nature Methods 2017;14:417-419.",
        "Muzellec B, Telenczuk M, Cabeli V, Andreux M. PyDESeq2: a Python package for bulk RNA-seq differential expression analysis. Bioinformatics 2023.",
        "Reiter JF, Leroux MR. Genes and molecular pathways underpinning ciliopathies. Nature Reviews Molecular Cell Biology 2017;18:533-547.",
        "NCBI BioProject PRJNA1513000 - RNA-seq of primary AML and KG1a cells with LNA/ZIP anti-miR-551b.",
    ]),
]

CSS = """
*{box-sizing:border-box} body{margin:0}
.paper{background:#fff;color:#1a1d21;max-width:800px;margin:0 auto;padding:34px 26px 70px;
  font-family:Georgia,"Times New Roman",serif;font-size:16.5px;line-height:1.62}
.paper h1{font-family:-apple-system,"Segoe UI",Roboto,sans-serif;font-size:1.55rem;line-height:1.25;
  text-align:center;margin:.2em 0 .4em;text-wrap:balance}
.byline{text-align:center;color:#5b6570;font-family:-apple-system,"Segoe UI",sans-serif;font-size:.95rem}
.affil{text-align:center;color:#5b6570;font-style:italic;font-size:.85rem;margin-bottom:20px}
.tag{text-align:center;font-family:-apple-system,"Segoe UI",sans-serif;font-size:.72rem;letter-spacing:.08em;
  text-transform:uppercase;color:#8a939d;margin-bottom:18px}
.abstract{background:#f5f7f9;border:1px solid #d7dce1;border-radius:8px;padding:16px 20px;font-size:.95rem;margin:0 0 26px}
.paper h2{font-family:-apple-system,"Segoe UI",Roboto,sans-serif;font-size:1.12rem;margin:30px 0 8px;
  padding-bottom:5px;border-bottom:2px solid #6a1b2a;color:#6a1b2a}
.paper h3{font-family:-apple-system,"Segoe UI",sans-serif;font-size:.98rem;margin:18px 0 4px;color:#7a2230}
.paper p{margin:.5em 0;text-align:justify}
figure{margin:20px 0;text-align:center} figure img{max-width:100%;border:1px solid #d7dce1;border-radius:6px}
figcaption{font-family:-apple-system,"Segoe UI",sans-serif;font-size:.8rem;color:#5b6570;margin-top:7px;text-align:left}
.wrap-x{overflow-x:auto;margin:16px 0}
table{border-collapse:collapse;width:100%;font-family:-apple-system,"Segoe UI",sans-serif;font-size:.82rem}
caption{font-family:-apple-system,"Segoe UI",sans-serif;font-size:.8rem;color:#5b6570;text-align:left;margin-bottom:6px}
th,td{border-bottom:1px solid #d7dce1;padding:6px 10px;text-align:left}
th{color:#6a1b2a;border-bottom:2px solid #6a1b2a} td:nth-child(n+2){text-align:right;font-variant-numeric:tabular-nums}
.refs{font-size:.82rem;color:#33424f} .refs ol{padding-left:20px} .refs li{margin:5px 0}
.note{background:#fff7ed;border-left:3px solid #e0872b;padding:10px 14px;font-size:.86rem;margin:16px 0;border-radius:0 6px 6px 0}
"""


def render_html(path):
    def uri(p):
        return "data:image/png;base64," + base64.b64encode(open(p, "rb").read()).decode()
    e = _html.escape
    out = ["<title>Cilia + AML (miR-551b) re-analysis</title>", f"<style>{CSS}</style>",
           '<article class="paper">', '<div class="tag">Secondary analysis · Transcriptomics</div>',
           f"<h1>{e(TITLE)}</h1>", f'<div class="byline">{e(BYLINE)}</div>',
           f'<div class="affil">{e(AFFIL)}</div>', f'<div class="abstract"><b>Abstract.</b> {e(ABSTRACT)}</div>']
    for b in BLOCKS:
        k = b[0]
        if k == "h2": out.append(f"<h2>{e(b[1])}</h2>")
        elif k == "h3": out.append(f"<h3>{e(b[1])}</h3>")
        elif k == "p": out.append(f"<p>{e(b[1])}</p>")
        elif k == "note": out.append(f'<div class="note">{e(b[1])}</div>')
        elif k == "fig": out.append(f'<figure><img src="{uri(b[1])}" alt=""><figcaption>{e(b[2])}</figcaption></figure>')
        elif k == "table":
            _, capt, headers, rows = b
            th = "".join(f"<th>{e(h)}</th>" for h in headers)
            body = "".join("<tr>" + "".join(f"<td>{e(str(c))}</td>" for c in r) + "</tr>" for r in rows)
            out.append(f'<div class="wrap-x"><table><caption>{e(capt)}</caption>'
                       f"<thead><tr>{th}</tr></thead><tbody>{body}</tbody></table></div>")
        elif k == "refs":
            out.append('<div class="refs"><ol>' + "".join(f"<li>{e(r)}</li>" for r in b[1]) + "</ol></div>")
    out.append("</article>")
    open(path, "w", encoding="utf-8").write("\n".join(out))


def render_docx(path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    doc = Document()
    doc.add_heading(TITLE, level=0)
    for txt, sz, ital in [(BYLINE, 11, False), (AFFIL, 9, True)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.size = Pt(sz); r.italic = ital
    ab = doc.add_paragraph(); ab.add_run("Abstract. ").bold = True; ab.add_run(ABSTRACT)
    for b in BLOCKS:
        k = b[0]
        if k == "h2": doc.add_heading(b[1], level=1)
        elif k == "h3": doc.add_heading(b[1], level=2)
        elif k in ("p", "note"):
            p = doc.add_paragraph(); run = p.add_run(b[1]); run.italic = (k == "note")
        elif k == "fig":
            doc.add_picture(b[1], width=Inches(5.7))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c = doc.add_paragraph().add_run(b[2]); c.italic = True; c.font.size = Pt(8.5)
        elif k == "table":
            _, capt, headers, rows = b
            c = doc.add_paragraph().add_run(capt); c.italic = True; c.font.size = Pt(9)
            tbl = doc.add_table(rows=1, cols=len(headers)); tbl.style = "Light Grid Accent 1"
            for i, h in enumerate(headers): tbl.rows[0].cells[i].text = h
            for row in rows:
                cells = tbl.add_row().cells
                for i, v in enumerate(row): cells[i].text = str(v)
        elif k == "refs":
            for i, ref in enumerate(b[1], 1):
                p = doc.add_paragraph(f"{i}. {ref}"); p.runs[0].font.size = Pt(9)
    doc.save(path)


def render_pdf(path):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle)
    def esc(s): return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    ss = getSampleStyleSheet()
    navy = colors.HexColor("#6a1b2a"); muted = colors.HexColor("#5b6570")
    h1 = ParagraphStyle("t", parent=ss["Title"], fontSize=13.5, leading=17, alignment=TA_CENTER)
    center = ParagraphStyle("c", parent=ss["Normal"], alignment=TA_CENTER, fontSize=9, textColor=muted)
    body = ParagraphStyle("b", parent=ss["Normal"], alignment=TA_JUSTIFY, fontSize=10, leading=14, spaceAfter=5)
    h2 = ParagraphStyle("h2", parent=ss["Heading2"], textColor=navy, fontSize=12, spaceBefore=10)
    h3 = ParagraphStyle("h3", parent=ss["Heading3"], fontSize=10.5, spaceBefore=6)
    cap = ParagraphStyle("cap", parent=ss["Normal"], fontSize=8, textColor=muted, spaceAfter=6)
    abst = ParagraphStyle("ab", parent=body, fontSize=9, leading=12, backColor=colors.HexColor("#f5f7f9"), borderPadding=8)
    note = ParagraphStyle("nt", parent=body, fontSize=9, backColor=colors.HexColor("#fff7ed"), borderPadding=6)
    refst = ParagraphStyle("rf", parent=body, fontSize=8.5, leading=11, spaceAfter=2)
    usable = A4[0] - 2 * inch
    story = [Paragraph(esc(TITLE), h1), Paragraph(esc(BYLINE), center), Paragraph(esc(AFFIL), center),
             Spacer(1, 10), Paragraph("<b>Abstract.</b> " + esc(ABSTRACT), abst), Spacer(1, 8)]
    for b in BLOCKS:
        k = b[0]
        if k == "h2": story.append(Paragraph(esc(b[1]), h2))
        elif k == "h3": story.append(Paragraph(esc(b[1]), h3))
        elif k == "p": story.append(Paragraph(esc(b[1]), body))
        elif k == "note": story.append(Paragraph("<b>" + esc(b[1]) + "</b>", note))
        elif k == "fig":
            img = Image(b[1]); img.drawWidth = usable; img.drawHeight = usable * img.imageHeight / img.imageWidth
            story += [Spacer(1, 4), img, Paragraph(esc(b[2]), cap)]
        elif k == "table":
            _, capt, headers, rows = b
            story += [Spacer(1, 4), Paragraph(esc(capt), cap)]
            ncol = len(headers)
            widths = [usable * 0.5] + [usable * 0.5 / (ncol - 1)] * (ncol - 1)
            tbl = Table([headers] + rows, hAlign="LEFT", colWidths=widths)
            tbl.setStyle(TableStyle([
                ("FONTSIZE", (0, 0), (-1, -1), 8.5), ("TEXTCOLOR", (0, 0), (-1, 0), navy),
                ("LINEBELOW", (0, 0), (-1, 0), 1, navy),
                ("LINEBELOW", (0, 1), (-1, -1), 0.3, colors.HexColor("#d7dce1")),
                ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
                ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
            story.append(tbl)
        elif k == "refs":
            for i, ref in enumerate(b[1], 1): story.append(Paragraph(f"{i}. {esc(ref)}", refst))
    SimpleDocTemplate(path, pagesize=A4, topMargin=0.9 * inch, bottomMargin=0.9 * inch,
                      leftMargin=inch, rightMargin=inch, title="Cilia+AML miR-551b re-analysis").build(story)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--html"); ap.add_argument("--docx"); ap.add_argument("--pdf")
    args = ap.parse_args()
    if args.html: render_html(args.html); print(f"[html] wrote {args.html}")
    if args.docx: render_docx(args.docx); print(f"[docx] wrote {args.docx}")
    if args.pdf: render_pdf(args.pdf); print(f"[pdf]  wrote {args.pdf}")


if __name__ == "__main__":
    main()
